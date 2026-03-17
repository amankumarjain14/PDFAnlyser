import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_heading_style(paragraph, level: int):
    """Apply colour and sizing to headings."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
    if level == 1:
        run.font.size = Pt(22)
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.bold = True
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        run.bold = True
    elif level == 3:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x37, 0x47, 0x51)
        run.bold = True


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _parse_table(doc: Document, lines: list, start_idx: int) -> int:
    """Parse a markdown table and add it to the document. Returns next line index."""
    table_lines = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            i += 1
        else:
            break

    if len(table_lines) < 2:
        return i

    # Remove separator row (e.g., |---|---|)
    rows = [l for l in table_lines if not re.match(r"^\|[-| :]+\|$", l)]
    if not rows:
        return i

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return i

    col_count = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(parsed):
        row_cells = table.rows[row_idx].cells
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < col_count:
                cell = row_cells[col_idx]
                cell.text = cell_text
                if row_idx == 0:
                    for run in cell.paragraphs[0].runs:
                        run.bold = True

    doc.add_paragraph()
    return i


def generate_docx(
    enhanced_content: str,
    original_filename: str,
    processed_date: str,
    word_count: int,
    output_path: str,
) -> str:
    """
    Convert markdown-formatted enhanced content into a styled .docx file.
    Returns the output_path.
    """
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Cover header ──────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("Enhanced Document Analysis")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    run.bold = True

    _add_horizontal_rule(doc)

    # Metadata block
    meta = doc.add_paragraph()
    meta.add_run("Original File: ").bold = True
    meta.add_run(original_filename)
    meta.add_run("\nProcessed: ").bold = True
    meta.add_run(processed_date)
    meta.add_run("\nVersion: ").bold = True
    meta.add_run("Enhanced v1.0")
    meta.add_run("\nWord Count: ").bold = True
    meta.add_run(str(word_count))

    _add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── Render markdown content ───────────────────────────────────────────────
    lines = enhanced_content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Heading levels
        if stripped.startswith("#### "):
            p = doc.add_heading(stripped[5:], level=3)
            _set_heading_style(p, 3)

        elif stripped.startswith("### "):
            p = doc.add_heading(stripped[4:], level=3)
            _set_heading_style(p, 3)

        elif stripped.startswith("## "):
            _add_horizontal_rule(doc)
            p = doc.add_heading(stripped[3:], level=2)
            _set_heading_style(p, 2)

        elif stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            _set_heading_style(p, 1)

        # Table detection
        elif stripped.startswith("|") and stripped.endswith("|"):
            i = _parse_table(doc, lines, i)
            continue

        # Bullet / list items
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:]
            p = doc.add_paragraph(style="List Bullet")
            _render_inline(p, text)

        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            p = doc.add_paragraph(style="List Number")
            _render_inline(p, text)

        # Blank line
        elif stripped == "" or stripped == "---":
            doc.add_paragraph()

        # Bold line (standalone)
        elif stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:-2])
            run.bold = True

        # Regular paragraph
        else:
            if stripped:
                p = doc.add_paragraph()
                _render_inline(p, stripped)

        i += 1

    # ── Footer ────────────────────────────────────────────────────────────────
    _add_horizontal_rule(doc)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_para.add_run(
        f"Generated by AI-Powered PDF Analysis | {datetime.now().strftime('%B %d, %Y')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run.italic = True

    doc.save(output_path)
    return output_path


def _render_inline(paragraph, text: str):
    """Render inline markdown bold (**text**) into paragraph runs."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
